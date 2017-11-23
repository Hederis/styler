import csv
import string
import glob
import os
import sys
import argparse
from distutils.util import strtobool
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

"""
stylenames.csv has one row for each *base* style name, and columns for the following:
prefix:
    "sec" = section-start style
    "wpr" = wrapper style
    "blk" = block style (paragraph)
    "inl" = inline style (character)
name: descriptive name for the style
levels: are levels allowed for this style? [boolean]
variations: are variations allowed? [boolean]
type: "1" for Paragraph, "2" for Character
order: does style need "first", "only", "last" versions? [boolean]

All versions for wrappers, levels, variations, etc. will be expanded automatically here; edit first variables below to adjust number of variations and number of levels created for each.

See README for explanation of style name syntax rules.
"""

def style_names(source_file, number_of_variations, number_of_levels, id_prefix):
    variations = list(string.ascii_lowercase[:number_of_variations])
    # range end is exclusive, so add 1
    levels = list(range(1,number_of_levels+1))
    wrappers = ["start", "end"]
    orders = ["first", "last", "only"]
    # document = Document()
    # docstyles = document.styles
    finalnames = []
    with open(source_file, 'rU') as csvfile:
        ourstyles = csv.DictReader(csvfile)
        for style in ourstyles:
            if int(style["type"]) == WD_STYLE_TYPE.PARAGRAPH:
                stylenames = [style["name"]]
                if strtobool(style["levels"]) == True:
                    stylenames = [x + str(y) for x in stylenames for y in levels]
                if strtobool(style["variations"]) == True:
                    stylenames = [x + "_" + y for x in stylenames for y in variations]
                if style["prefix"] == "wpr":
                    stylenames = [x + "_" + y for x in stylenames for y in wrappers]
                if strtobool(style["order"]) == True:
                    stylenames = [x + "_" + y for x in stylenames for y in orders] + [style["name"]]

                finalnames = finalnames + ["_".join([id_prefix,style["prefix"],x]) for x in stylenames]
    return finalnames

def add_styles_to_doc(word_docx, list_of_styles):
    document = Document(word_docx)
    docstyles = document.styles
    for stylename in list_of_styles:
        if stylename not in docstyles:
            docstyles.add_style(stylename, WD_STYLE_TYPE.PARAGRAPH)
            # p = document.add_paragraph(finalname, style=finalname)
    document.save(word_docx)

# program arguments
parser = argparse.ArgumentParser(description="Adds our stylenames to a Word file.")
parser.add_argument("input_doc", help="Path to docx file or directory containing docx files.")
parser.add_argument("--variations", default=1, type=int, dest="max_variations", help="Specify the maximum number of variations created for each style that allows variations. Default is 1.")
parser.add_argument("--levels", default=1, type=int, dest="max_levels", help="Specify the maximum number of levels created for each style that allows levels. Default is 1.")
parser.add_argument("--style-data", default="stylenames.csv", dest="style_source_file", help="Name of CSV file containing data about the style names to be added.")
parser.add_argument("--style-id", default="h", dest="identifier", help="Single character used to distinguish our styles from others that may be present in the document.")

args = parser.parse_args()

# generate list of style names to add
styles = style_names(args.style_source_file, args.max_variations, args.max_levels, args.identifier)

# process doc or docs
docs = args.input_doc
if os.path.exists(docs):
    if os.path.isdir(docs):
        os.chdir(docs)
        # reverse to print countdown as files are processed
        for i, file in enumerate(glob.glob("*.docx")):
            print i
            add_styles_to_doc(file, styles)
    elif os.path.isfile(docs):
        add_styles_to_doc(docs, styles)
